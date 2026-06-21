from __future__ import annotations

import json
import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"D:\BUID\Year 4 2025 - 2026\third semester\Special Topics in AI\Project\climate_evidence_graphrag_agent")
OUT = ROOT / "deliverables" / "EXPLANATION_GUIDES"
OUT.mkdir(parents=True, exist_ok=True)
GENERATED = datetime.now().strftime("%Y-%m-%d %H:%M")


MEMBERS = {
    "Reem": {
        "role": "Data ingestion, corpus quality, page and citation verification",
        "d1": ["reports/member1_d1_data_section.md", "scripts/build_d1_proxy_eval_set.py"],
        "notebooks": [
            "notebooks/D2_01_Reem_ingestion_data_quality.ipynb",
            "notebooks/D3_01_Reem_page_citation_verification.ipynb",
        ],
        "support": [
            "src/ingest/pdf_loader.py",
            "src/ingest/chunker.py",
            "src/ingest/page_verifier.py",
            "reports/tables/d2_ingestion_summary.csv",
            "reports/tables/page_citation_check.csv",
        ],
        "say": "I prepared and audited the evidence base. My part makes sure PDFs become searchable chunks with document IDs and page numbers, so the final GraphRAG answer can defend where evidence came from.",
        "qs": [
            ("Where does the text come from?", "From the PDFs and metadata in the project data folders. The ingestion code reads PDF pages, creates smaller chunks, and keeps page numbers."),
            ("Why do we chunk documents?", "The model and retriever cannot use a whole PDF at once. Chunks are smaller cards that can be searched."),
            ("Why overlap chunks?", "Overlap prevents an important sentence from being cut in half between two chunks."),
            ("How do you know citations are real?", "The page-citation notebook checks document ID and page fields and saves reports/tables/page_citation_check.csv."),
            ("What could go wrong here?", "PDF text extraction can be messy, metadata can be missing, chunks can be duplicated, or pages can be wrong."),
        ],
    },
    "Salma": {
        "role": "Retrieval, AutoML, hybrid search, ablation, and latency",
        "d1": [
            "reports/member2_d1_automl_section.md",
            "src/retrieval/automl_tuner.py",
            "reports/tables/d1_baseline_vs_automl_metrics.csv",
            "reports/tables/d1_automl_trials.csv",
        ],
        "notebooks": [
            "notebooks/D2_02_Salma_retrieval_comparison.ipynb",
            "notebooks/D3_02_Salma_retrieval_ablation.ipynb",
            "notebooks/D4_02_Salma_retrieval_latency.ipynb",
        ],
        "support": [
            "src/retrieval/bm25_retriever.py",
            "src/retrieval/dense_retriever.py",
            "src/retrieval/hybrid_retriever.py",
            "src/retrieval/fusion.py",
            "reports/tables/d2_retrieval_metrics_summary.csv",
            "reports/tables/d3_retrieval_ablation_summary.csv",
            "reports/tables/d4_retrieval_latency_summary.csv",
        ],
        "say": "I evaluated retrieval. The system combines BM25 keyword search with dense semantic search, then compares retrieval settings and latency so we know why hybrid retrieval is used.",
        "qs": [
            ("What is BM25?", "A keyword search method. It rewards chunks that contain the same words as the query."),
            ("What is dense retrieval?", "A meaning-based search. Text is changed into numbers, and similar meanings are matched."),
            ("Why hybrid retrieval?", "BM25 is good for exact words; dense search is good for meaning. Hybrid uses both."),
            ("What did AutoML tune?", "It searched settings such as top-k, metric, SVD dimension, normalization, and BM25/dense weight."),
            ("Why include D4 latency?", "D4 was merged with final D3 work, but the latency notebook still proves the app can respond in an acceptable time."),
        ],
    },
    "Rana": {
        "role": "Knowledge graph, Neo4j, Cypher templates, and GraphRAG executor",
        "d1": ["reports/member3_d1_graph_plan_section.md", "src/graph/graph_schema.md"],
        "notebooks": [
            "notebooks/D2_03_Rana_graph_build_cypher.ipynb",
            "notebooks/D3_03_Rana_graphrag_executor.ipynb",
        ],
        "support": [
            "src/graph/neo4j_builder.py",
            "src/graph/cypher_queries.py",
            "src/rag/graphrag_executor.py",
            "reports/tables/d2_graph_counts.csv",
            "reports/tables/d3_graph_guided_results.csv",
        ],
        "say": "I built and tested the graph side. The graph stores climate evidence relationships in Neo4j. Python chooses a Cypher template, sends it to Neo4j, and uses returned graph evidence in GraphRAG.",
        "qs": [
            ("What language is the graph query written in?", "Cypher. Cypher is Neo4j's graph query language. Python only sends the Cypher string."),
            ("Where is Cypher created?", "The prepared Cypher templates are in src/graph/cypher_queries.py."),
            ("Where is Cypher executed?", "In src/rag/graphrag_executor.py through the Neo4j driver session.run call."),
            ("Why use a graph?", "A graph stores explicit links such as Document supports Finding or Country has Policy."),
            ("What if the graph has no answer?", "The system falls back to hybrid BM25 plus dense retrieval."),
        ],
    },
    "Aaya": {
        "role": "Online learning, adaptive retrieval, final demo, and PEFT/QLoRA tuning",
        "d1": [
            "reports/member4_d1_online_learning_section.docx",
            "src/learning/river_topic_classifier.py",
            "src/learning/feedback_adapter.py",
            "src/learning/drift_detector.py",
            "reports/tables/prequential_online_learning_results.csv",
        ],
        "notebooks": [
            "notebooks/D2_04_Aaya_online_retrieval_adaptation.ipynb",
            "notebooks/D3_04_Aaya_online_graphrag_adaptation.ipynb",
            "notebooks/D3_06_Final_Demo_Tuning_Merged_Scope.ipynb",
            "notebooks/D3_07_Kaggle_QLoRA_Tuning.ipynb",
        ],
        "support": [
            "app/streamlit_app.py",
            "src/learning/feedback_adapter.py",
            "src/learning/river_topic_classifier.py",
            "src/learning/drift_detector.py",
            "data/tuning/finetune_qa.jsonl",
            "models/qlora_adapter/adapter_config.json",
            "reports/tables/d3_online_retrieval_comparison.csv",
            "reports/tables/d3_or_final_zero_shot_vs_tuned.csv",
            "reports/tables/d3_tuning_latency.csv",
        ],
        "say": "I handled online adaptation, final demo behavior, and tuning. The online learner updates from feedback, and QLoRA trains a small adapter instead of the full model.",
        "qs": [
            ("What is online learning?", "Learning step by step as examples arrive, instead of training only once."),
            ("What is prequential evaluation?", "Predict first, then learn from the correct answer. This simulates a real stream."),
            ("How does feedback change retrieval?", "Helpful feedback rewards a weight choice; unhelpful feedback moves the adapter away from that choice."),
            ("Why can adaptive retrieval look worse?", "If feedback labels are noisy, the adapter may learn from wrong signals."),
            ("What is QLoRA?", "A memory-saving fine-tuning method that trains a small adapter on a quantized model."),
            ("Where is the user query created in the demo?", "In app/streamlit_app.py. Streamlit creates the query from the example buttons or from st.text_input, then passes it to executor.run(query)."),
        ],
    },
    "Alia": {
        "role": "API, integration, RAG quality, safety, and citation evaluation",
        "d1": ["reports/member5_d1_eval_section.md", "src/evaluation/retrieval_metrics.py", "src/evaluation/latency.py"],
        "notebooks": [
            "notebooks/D2_05_Alia_api_tests_integration.ipynb",
            "notebooks/D3_05_Alia_safety_rag_evaluation.ipynb",
            "notebooks/D3_graphrag_eval_safety.ipynb",
        ],
        "support": [
            "src/api/main.py",
            "src/evaluation/rag_metrics.py",
            "src/safety/citation_verifier.py",
            "src/safety/source_pinning.py",
            "reports/tables/d3_rag_eval_metrics.csv",
            "reports/tables/d3_safety_before_after.csv",
            "reports/tables/d3_app_query_validation_no_gemini.csv",
        ],
        "say": "I checked integration and safety. My part verifies API behavior, answer quality, citation correctness, fallback behavior, and app-facing evaluation outputs.",
        "qs": [
            ("Where is the search API calling?", "In src/api/main.py, the POST /search endpoint calls retriever.search(req.question, ...)."),
            ("What is the difference between FastAPI and Streamlit?", "FastAPI is a programmatic web API. Streamlit is the visual demo app."),
            ("How is answer quality measured?", "Using faithfulness, answer relevance, citation correctness, and latency tables."),
            ("What is citation safety?", "The answer should cite retrieved evidence and avoid unsupported claims."),
            ("What if Gemini fails?", "Fallback behavior still shows retrieved evidence and validated saved outputs can be used."),
        ],
    },
}


GLOSSARY = [
    ("Notebook", "A .ipynb workbook with text cells and code cells."),
    ("Cell", "One block inside a notebook. Code cells run; markdown cells explain."),
    ("Input", "What the code reads first: a question, CSV, JSON, PDF, config, or folder."),
    ("Output", "What the code creates: a table, metric, chart, CSV, answer, or saved model."),
    ("DataFrame", "A Python table, similar to an Excel sheet in memory."),
    ("CSV", "A simple table file with commas between values."),
    ("JSON/JSONL", "Structured text. JSONL stores one record per line."),
    ("BM25", "Keyword retrieval that likes exact word matches."),
    ("Dense retrieval", "Meaning retrieval where text becomes numbers and similar meanings match."),
    ("Hybrid retrieval", "BM25 plus dense retrieval together."),
    ("SVD", "A compression method that reduces many text features into fewer dimensions."),
    ("Neo4j", "The graph database used for entities and relationships."),
    ("Cypher", "The query language used by Neo4j."),
    ("RAG", "Retrieve evidence first, then generate an answer from that evidence."),
    ("GraphRAG", "RAG plus graph reasoning from Neo4j."),
    ("Gemini", "The language model used to write final answers."),
    ("River", "A Python library for online learning one example at a time."),
    ("ADWIN", "A drift detector that watches for changing feedback/performance."),
    ("Prequential", "Predict first, then learn. Used for stream evaluation."),
    ("QLoRA", "Efficient fine-tuning that trains a small adapter on a quantized model."),
]


REPL = {
    "\u2013": "-",
    "\u2014": "-",
    "\u2018": "'",
    "\u2019": "'",
    "\u201c": '"',
    "\u201d": '"',
    "\u2192": "->",
    "\u21d2": "=>",
    "\u2022": "-",
    "\u2026": "...",
    "\u00a0": " ",
    "\u2265": ">=",
    "\u2264": "<=",
    "\u2248": "~",
    "\ufeff": "",
    "\u00b0": " degrees ",
}


def clean(x) -> str:
    if x is None:
        return ""
    x = str(x)
    for a, b in REPL.items():
        x = x.replace(a, b)
    x = x.encode("ascii", "replace").decode("ascii")
    # WordprocessingML cannot contain null bytes or most ASCII control chars.
    # Keep normal tab/newline/carriage-return, remove the rest.
    x = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", x)
    return x


def rpath(p) -> str:
    p = Path(p)
    try:
        return str(p.relative_to(ROOT))
    except Exception:
        return str(p)


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def text_cell(cell, text, bold=False, size=8.5, font="Calibri"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(clean(text))
    run.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def table(doc, rows, headers=None, widths=None, size=8.5):
    all_rows = ([headers] if headers else []) + rows
    if not all_rows:
        return None
    t = doc.add_table(rows=len(all_rows), cols=len(all_rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(all_rows):
        for j, val in enumerate(row):
            c = t.cell(i, j)
            text_cell(c, val, bold=bool(headers and i == 0), size=size)
            if headers and i == 0:
                shade(c, "E8EEF5")
    if widths:
        for rr in t.rows:
            for j, w in enumerate(widths):
                rr.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def setup(doc: Document, title: str):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
    for nm, font, size, after in [("CodeSmall", "Consolas", 7.5, 1), ("ExplainSmall", "Calibri", 9, 3)]:
        try:
            st = styles.add_style(nm, 1)
        except Exception:
            st = styles[nm]
        st.font.name = font
        st._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        st.font.size = Pt(size)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.0
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(clean(title))
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Climate Evidence GraphRAG Agent - generated " + GENERATED)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("555555")


def para(doc, text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.add_run(clean(text))
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(clean(it))


def nums(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(clean(it))


def read_file(path, limit=5000):
    path = Path(path)
    if not path.exists():
        return "Missing file: " + rpath(path)
    if path.suffix.lower() == ".docx":
        try:
            d = Document(path)
            return "\n".join(p.text for p in d.paragraphs if p.text.strip())[:limit]
        except Exception as e:
            return "Could not read DOCX: " + str(e)
    return path.read_text(encoding="utf-8", errors="replace")[:limit]


def explain(line: str) -> str:
    s = line.strip()
    if not s:
        return "Blank line for readability."
    if s.startswith("#"):
        return "A human comment. Python ignores it; it explains the code or marks a section."
    if s.startswith("!") or s.startswith("%"):
        return "A notebook shortcut command. It controls the notebook environment, not project logic."
    if s.startswith("import "):
        return "Loads a Python library so this notebook can use its ready-made tools."
    m = re.match(r"from\s+([\w\.]+)\s+import\s+(.+)", s)
    if m:
        return f"Loads {m.group(2)} from {m.group(1)} so this notebook can call that code."
    if s.startswith("def "):
        return "Defines a reusable function. A function is a small machine: inputs go in, work happens, output comes back."
    if s.startswith("class "):
        return "Defines a class, which is a blueprint for an object that keeps data and actions together."
    if s.startswith("return "):
        return "Sends a value back from a function to the code that called it."
    if s in ("try:", "try :"):
        return "Starts a safe attempt. If this risky code fails, the except block can handle it."
    if s.startswith("except"):
        return "Handles an error so the notebook can show a controlled message or continue."
    if s.startswith("finally"):
        return "Runs cleanup code whether the earlier step succeeded or failed."
    if s.startswith("if "):
        return "Checks a condition. If true, Python runs the indented block under it."
    if s.startswith("elif "):
        return "Checks another condition if the earlier condition was false."
    if s.startswith("else"):
        return "Runs this branch when the earlier checks were false."
    if s.startswith("for "):
        return "Starts a loop. Python repeats the indented code for every item in a list/table/range."
    if s.startswith("while "):
        return "Starts a loop that keeps running while the condition stays true."
    if s.startswith("with "):
        return "Opens a resource safely, such as a file or database session, and closes it after use."
    if "pd.read_csv" in s:
        return "Reads a CSV file into a pandas DataFrame, which is like an Excel table in Python."
    if "json.load" in s or "json.loads" in s or "pd.read_json" in s:
        return "Reads JSON structured text and turns it into Python data."
    if ".to_csv" in s:
        return "Saves a DataFrame as a CSV proof/output file."
    if ".to_json" in s or ".write_text" in s:
        return "Writes data or text to a saved file."
    if "display(" in s:
        return "Shows a table/object nicely inside the notebook output."
    if s.startswith("print("):
        return "Prints a value or progress message for the person running the notebook."
    if "Path(" in s or "PROJECT_ROOT" in s or "ROOT" in s:
        return "Creates or uses a project file path so the notebook knows where files are."
    if "GraphRAGExecutor.from_config" in s:
        return "Builds the full GraphRAG pipeline from configs/config.yaml."
    if "executor.run" in s or "invoke_graphrag" in s:
        return "Sends a question into GraphRAG and receives evidence, answer, citations, and trace data."
    if "GraphDatabase.driver" in s:
        return "Creates the Neo4j connection object used by Python to talk to the graph database."
    if "session.run" in s:
        return "Executes a Cypher query in Neo4j and returns graph rows."
    if "retriever.search" in s or "hybrid.search" in s or ".search(" in s:
        return "Runs a retrieval search: the question goes in and ranked chunks come out."
    if "BM25" in s or "bm25" in s:
        return "Uses/configures BM25 keyword search."
    if "dense" in s.lower() or "svd" in s.lower() or "TfidfVectorizer" in s:
        return "Uses semantic/numeric retrieval features, where text becomes numbers for similarity search."
    if "River" in s or "river" in s or "learn_one" in s or "predict_one" in s:
        return "Uses online learning: predict one example, then learn from it."
    if "ADWIN" in s or "drift" in s.lower():
        return "Uses drift detection to notice when feedback/performance changes over time."
    if "SFTTrainer" in s or "LoraConfig" in s or "BitsAndBytesConfig" in s or "prepare_model_for_kbit_training" in s:
        return "Part of QLoRA/PEFT tuning. It trains a small adapter while saving GPU memory."
    if "st." in s:
        return "A Streamlit command that creates something visible in the web app."
    if "@app." in s or "FastAPI" in s:
        return "Defines or uses a FastAPI web route for programmatic access."
    if "=" in s and not s.startswith(("==", ">=", "<=", "!=")):
        left = s.split("=", 1)[0].strip()
        if len(left) < 55:
            return "Stores the value on the right side into the variable " + left + ". Later code can reuse it."
        return "Stores a computed result in a variable for later use."
    if s.endswith(")"):
        return "Calls a function, which means it asks a named tool to do work using the values in parentheses."
    return "Runs one Python instruction. Read it as: calculate values, pass them to functions, or control what happens next."


def md_title(src):
    for line in src.splitlines():
        t = line.strip().lstrip("#").strip()
        if t:
            return t[:120]
    return "Markdown context"


def out_summary(outputs):
    if not outputs:
        return "No saved output in this cell, or output was cleared."
    parts = []
    for o in outputs[:2]:
        typ = o.get("output_type")
        if typ == "stream":
            txt = "".join(o.get("text", ""))
            txt = re.sub(r"\s+", " ", txt).strip()
            if txt:
                parts.append("Printed: " + txt[:180])
        elif typ in ("execute_result", "display_data"):
            data = o.get("data", {})
            if "text/plain" in data:
                txt = data["text/plain"]
                txt = "".join(txt) if isinstance(txt, list) else str(txt)
                txt = re.sub(r"\s+", " ", txt).strip()
                parts.append("Displayed: " + txt[:180])
            elif "image/png" in data:
                parts.append("Displayed a chart/image.")
            else:
                parts.append("Displayed notebook output.")
        elif typ == "error":
            parts.append("Saved error output: " + o.get("ename", "error") + " - " + o.get("evalue", ""))
    return " | ".join(parts) if parts else "Output exists but is not text-based."


def file_mentions(src):
    ms = []
    for m in re.finditer(r"['\"]([^'\"]+\.(csv|jsonl|json|yaml|yml|pkl|png|docx|pdf|md|parquet|txt))['\"]", src):
        ms.append(m.group(1))
    return sorted(set(ms))[:20]


def glossary(doc):
    doc.add_heading("Simple dictionary", level=1)
    table(doc, GLOSSARY, headers=["Word", "Meaning"], widths=[1.5, 4.8], size=8.5)


def add_py_lines(doc, path, max_lines=None):
    lines = Path(path).read_text(encoding="utf-8", errors="replace").splitlines()
    count = 0
    if max_lines:
        para(doc, "To keep this guide usable, this source-module section explains the first important non-empty lines. Notebook cells are explained fully.", style="ExplainSmall")
    for i, line in enumerate(lines, 1):
        if not line.strip():
            continue
        count += 1
        if max_lines and count > max_lines:
            para(doc, "Remaining shared source code is summarized in the master technical guide.", style="ExplainSmall")
            break
        cp = doc.add_paragraph(style="CodeSmall")
        cp.add_run(clean(f"Line {i}: {line}"))
        ep = doc.add_paragraph(style="ExplainSmall")
        ep.add_run("Kid-simple meaning: ").bold = True
        ep.add_run(clean(explain(line)))


def source_section(doc, title, paths):
    doc.add_heading(title, level=1)
    for ps in paths:
        p = ROOT / ps
        doc.add_heading(rpath(p), level=2)
        para(doc, "File exists: " + str(p.exists()))
        if not p.exists():
            continue
        ext = p.suffix.lower()
        if ext in [".md", ".txt", ".csv", ".json", ".jsonl", ".yaml", ".yml"]:
            para(doc, "What it is: written evidence, config, data, or output artifact connected to this deliverable.")
            cp = doc.add_paragraph(style="CodeSmall")
            cp.add_run(clean(read_file(p, 3500)))
        elif ext == ".docx":
            para(doc, "What it is: Word report section connected to this deliverable.")
            cp = doc.add_paragraph(style="ExplainSmall")
            cp.add_run(clean(read_file(p, 3500)))
        elif ext == ".py":
            para(doc, "What it is: Python source code. Important active lines are explained below.")
            add_py_lines(doc, p, max_lines=260)
        else:
            para(doc, "Supporting artifact for this deliverable.")


def add_notebook(doc, nb_path):
    nb_path = Path(nb_path)
    doc.add_heading(rpath(nb_path), level=1)
    if not nb_path.exists():
        para(doc, "Notebook missing.")
        return
    nb = json.loads(nb_path.read_text(encoding="utf-8"))
    cells = nb.get("cells", [])
    code = [c for c in cells if c.get("cell_type") == "code"]
    md = [c for c in cells if c.get("cell_type") == "markdown"]
    allsrc = "\n".join("".join(c.get("source", [])) for c in cells)
    para(doc, f"Notebook summary: {len(cells)} total cells, {len(code)} code cells, {len(md)} markdown cells.")
    ms = file_mentions(allsrc)
    if ms:
        para(doc, "Likely inputs/outputs/config files mentioned by this notebook:")
        bullets(doc, ms)
    last = "Notebook start"
    ci = 0
    for idx, c in enumerate(cells, 1):
        src = "".join(c.get("source", []))
        typ = c.get("cell_type")
        if typ == "markdown":
            last = md_title(src)
            if src.strip() and (src.lstrip().startswith("#") or len(src.strip()) < 600):
                doc.add_heading("Markdown context: " + clean(last), level=3)
                para(doc, re.sub(r"\s+", " ", src.strip())[:700], style="ExplainSmall")
            continue
        if typ != "code":
            continue
        ci += 1
        doc.add_heading(f"Code cell {ci} - {clean(last)}", level=2)
        para(doc, f"Notebook cell number: {idx}. Output clue: {out_summary(c.get('outputs', []))}", style="ExplainSmall")
        lines = [(i, l) for i, l in enumerate(src.splitlines(), 1) if l.strip()]
        if not lines:
            para(doc, "This code cell is empty.")
            continue
        for ln, line in lines:
            cp = doc.add_paragraph(style="CodeSmall")
            cp.add_run(clean(f"Cell line {ln}: {line}"))
            ep = doc.add_paragraph(style="ExplainSmall")
            ep.add_run("Kid-simple meaning: ").bold = True
            ep.add_run(clean(explain(line)))


def member_doc(name, cfg):
    doc = Document()
    setup(doc, f"{name} Code Explanation Guide - D1, D2, D3/final")
    para(doc, "Member role: " + cfg["role"])
    para(doc, "This guide explains the member's code and artifacts in simple language. It includes inputs, outputs, systems used, technical meanings, code-line explanations, and likely doctor questions.")
    doc.add_heading("Files explained", level=1)
    bullets(doc, cfg["d1"] + cfg["notebooks"] + cfg["support"])
    glossary(doc)
    source_section(doc, "D1 files and supporting code", cfg["d1"])
    has_d4 = any("D4_" in nb for nb in cfg["notebooks"])
    heading = "D2, D3, and included D4 notebooks - line by line" if has_d4 else "D2 and D3 notebooks - line by line"
    doc.add_heading(heading, level=1)
    if has_d4:
        para(doc, "Every non-empty code line in the assigned notebooks is shown and explained. A D4 notebook exists for this member, so it is included because D4 work was merged into the final D3/final submission.")
    else:
        para(doc, "Every non-empty code line in the assigned notebooks is shown and explained.")
    for nb in cfg["notebooks"]:
        add_notebook(doc, ROOT / nb)
    doc.add_heading("Supporting files to know", level=1)
    for ps in cfg["support"]:
        p = ROOT / ps
        doc.add_heading(rpath(p), level=2)
        para(doc, "Exists: " + str(p.exists()))
        if p.exists() and p.suffix.lower() in [".csv", ".json", ".jsonl", ".yaml", ".yml", ".md", ".txt"]:
            cp = doc.add_paragraph(style="CodeSmall")
            cp.add_run(clean(read_file(p, 1600)))
        elif p.exists() and p.suffix.lower() == ".py":
            para(doc, "Shared Python module. The member should know what it does; exact project flow is covered in the master guide.")
    doc.add_heading("Expected doctor questions for this member", level=1)
    for q, a in cfg["qs"]:
        doc.add_heading("Q: " + q, level=3)
        para(doc, "A: " + a)
    doc.add_heading("One-minute speaking script", level=1)
    para(doc, cfg["say"])
    out = OUT / f"{name}_Code_Explanation_D1_D3.docx"
    doc.save(out)
    return out


def master_doc():
    doc = Document()
    setup(doc, "Full Project Technical Deep Explanation and Viva Q&A")
    para(doc, "This document explains the full implementation deeply enough for presentation questions. It focuses on how data moves through the system and where key code lives.")
    doc.add_heading("Project pipeline in simple terms", level=1)
    nums(doc, [
        "Read PDFs and metadata.",
        "Split PDF text into chunks with document IDs and page numbers.",
        "Build BM25 keyword and dense/SVD retrieval indexes.",
        "Build a Neo4j graph from documents, findings, topics, policies, countries, and targets.",
        "For a user question, run GraphRAG: classify intent, choose Cypher, query Neo4j, expand graph hits, blend with hybrid retrieval, call Gemini, return answer and citations.",
        "Show live trace in Streamlit and save evaluation evidence in reports/tables.",
    ])
    doc.add_heading("Where important things happen", level=1)
    rows = [
        ("User query is created", "app/streamlit_app.py", "Example buttons and st.text_input create the query variable as a Python string."),
        ("App sends query", "app/streamlit_app.py", "executor.run(query) calls GraphRAG."),
        ("Search API", "src/api/main.py", "POST /search receives JSON and calls retriever.search."),
        ("Hybrid retrieval", "src/retrieval/hybrid_retriever.py", "Calls BM25 and dense retrieval, then fuses results."),
        ("BM25", "src/retrieval/bm25_retriever.py", "Keyword search."),
        ("Dense/SVD", "src/retrieval/dense_retriever.py", "Meaning-style numeric retrieval."),
        ("Fusion", "src/retrieval/fusion.py", "Combines rankings/scores."),
        ("Cypher templates", "src/graph/cypher_queries.py", "Prepared Neo4j Cypher query strings."),
        ("Graph build", "src/graph/neo4j_builder.py", "Creates graph nodes and relationships."),
        ("GraphRAG", "src/rag/graphrag_executor.py", "Runs graph retrieval, fallback, blending, and Gemini generation."),
        ("Online learning", "src/learning/*.py", "River classifier, feedback adapter, drift detector."),
        ("Evaluation/safety", "src/evaluation/*.py and src/safety/*.py", "Metrics, citation checks, source pinning."),
    ]
    table(doc, rows, headers=["Part", "File", "Explanation"], widths=[1.5, 2.1, 2.8], size=8)
    doc.add_heading("Query path - exact explanation", level=1)
    bullets(doc, [
        "The user question starts as English text in the Streamlit input box.",
        "In code, that English text is stored as a Python string named query.",
        "When the Ask button is clicked, Streamlit calls executor.run(query).",
        "Inside GraphRAGExecutor.run, the system may classify intent/topic and choose a Cypher template.",
        "Cypher is the Neo4j graph query language. It is stored as text templates in src/graph/cypher_queries.py.",
        "Python sends the Cypher plus parameters to Neo4j using session.run.",
        "The graph rows become graph hits, then graph hits are expanded to text chunks and blended with hybrid retrieval chunks.",
        "The final evidence is placed into an English prompt and sent to Gemini for answer generation.",
    ])
    doc.add_heading("What language is each query/action written in", level=1)
    table(doc, [
        ("User question", "English natural language", "Typed in Streamlit or sent in API JSON."),
        ("Python call", "Python", "executor.run(query), retriever.search(query)."),
        ("Graph query", "Cypher", "MATCH patterns for Neo4j."),
        ("API request", "HTTP/JSON", "POST /search with question and top_k."),
        ("LLM prompt", "English text", "Evidence-grounded instruction sent to Gemini."),
    ], headers=["Thing", "Language/format", "Where"], widths=[1.6, 1.6, 3.0], size=8.5)
    doc.add_heading("D1 technical summary", level=1)
    bullets(doc, [
        "AutoML tuned retrieval settings and compared baseline vs tuned retrieval using NDCG@5, Recall@5, and p95 latency.",
        "Online learning used streaming/prequential evaluation: predict first, then learn.",
        "ADWIN drift detection was included to notice changing feedback/performance.",
        "Run-card configs and CSV outputs make the result reproducible.",
    ])
    doc.add_heading("D2 technical summary", level=1)
    table(doc, [
        ("Reem", "Ingestion and data quality", "PDF/chunk/metadata/page verification."),
        ("Salma", "Retrieval comparison", "BM25, dense, hybrid metrics and latency."),
        ("Rana", "Graph build", "Neo4j schema/counts/Cypher evidence."),
        ("Aaya", "Online adaptation", "Static vs adaptive retrieval and feedback effects."),
        ("Alia", "API/evaluation", "API tests and integration checks."),
    ], headers=["Member", "Area", "What it proves"], widths=[1.0, 1.8, 3.5], size=8.5)
    doc.add_heading("D3/final technical summary", level=1)
    bullets(doc, [
        "GraphRAG executor connects the graph, retriever, evidence blender, Gemini answer generator, and citations.",
        "The Streamlit app presents live GraphRAG trace plus a hybrid retrieval panel.",
        "D4 latency work is included as final performance evidence, especially Salma's retrieval latency notebook.",
        "QLoRA/PEFT is included through finetune_qa.jsonl, models/qlora_adapter, and comparison/latency tables.",
        "Safety evaluation checks faithfulness, answer relevance, citation correctness, and fallback behavior.",
    ])
    doc.add_heading("Likely doctor questions and answers", level=1)
    qs = [
        ("Where is the query created?", "In app/streamlit_app.py. st.text_input returns the user's typed English question as a Python string. Example buttons set the same input."),
        ("What language is the graph query written in?", "Cypher, because Neo4j uses Cypher. The user question is English; Python selects a Cypher template."),
        ("Where is Cypher stored?", "src/graph/cypher_queries.py."),
        ("Where is Cypher executed?", "src/rag/graphrag_executor.py through Neo4j driver session.run."),
        ("Where is the search API calling?", "src/api/main.py, POST /search calls retriever.search(req.question, k=req.top_k, filters=req.filters)."),
        ("Does Streamlit call FastAPI?", "For the local demo, Streamlit mainly calls project Python classes directly. FastAPI exists for external/programmatic access."),
        ("Why GraphRAG?", "It adds structured graph relationships to normal text retrieval, useful for entity/path questions."),
        ("When does fallback happen?", "If graph retrieval is empty or not useful, the system uses hybrid BM25 plus dense retrieval."),
        ("Where is Gemini called?", "In src/rag/graphrag_executor.py inside the answer generation part. The prompt contains query, graph hits, blended chunks, and citation instructions."),
        ("What is the search API output?", "JSON with query, top_k, filters, retriever description, and ranked result objects."),
        ("How do you prove citations?", "Chunks keep document ID/page number, and page-citation tables audit that these fields are present."),
        ("Why QLoRA instead of full fine-tuning?", "Full fine-tuning is too expensive. QLoRA trains only adapter weights on a quantized model."),
        ("What is the biggest limitation?", "The graph can only answer entities/relationships represented in the curated graph. Otherwise fallback is needed."),
    ]
    for q, a in qs:
        doc.add_heading("Q: " + q, level=3)
        para(doc, "A: " + a)
    doc.add_heading("Emergency demo fallback explanation", level=1)
    para(doc, "If Neo4j or Gemini fails live, say: the live demo depends on external Neo4j and Gemini services. The local retrieval and saved audited outputs still prove the system. The right panel shows hybrid retrieval fallback, and reports/tables/d3_graph_guided_results.csv contains the validated graph-guided run.")
    out = OUT / "Full_Project_Technical_Deep_Explanation_and_Viva_QA.docx"
    doc.save(out)
    return out


def demo_doc():
    doc = Document()
    setup(doc, "Team Streamlit Demo Speaking Script")
    para(doc, "Use this script when presenting the Streamlit app. It tells each member what to say and prepares answers for implementation questions.")
    doc.add_heading("Demo goal", level=1)
    para(doc, "Show the final app as an evidence-grounded GraphRAG system, not only a chatbot: PDF chunks, hybrid retrieval, Neo4j graph, fallback, Gemini answer, and citations.")
    doc.add_heading("Before demo checklist", level=1)
    nums(doc, [
        "Open the main project folder.",
        "Activate the .venv if needed.",
        "Check .env has Neo4j and Gemini credentials if using live graph/generation.",
        "Run streamlit run app/streamlit_app.py.",
        "Keep reports/tables/d3_graph_guided_results.csv ready as validated backup.",
    ])
    doc.add_heading("Speaking order", level=1)
    table(doc, [
        ("0:00-0:45", "Aaya", "Introduce app and full pipeline."),
        ("0:45-1:30", "Reem", "Explain PDFs, chunks, metadata, page citations."),
        ("1:30-2:20", "Salma", "Explain BM25, dense, hybrid retrieval, metrics, latency."),
        ("2:20-3:20", "Rana", "Explain Neo4j graph, Cypher templates, graph execution."),
        ("3:20-4:10", "Aaya", "Explain online adaptation and QLoRA."),
        ("4:10-5:00", "Alia", "Explain API, safety, citation evaluation, fallback."),
        ("5:00-7:00", "Team", "Run app questions and answer doctor questions."),
    ], headers=["Time", "Speaker", "What to say/show"], widths=[1.0, 1.0, 4.2], size=8.5)
    scripts = {
        "Aaya opening": "Our final demo is a climate evidence GraphRAG agent. The user asks a question in Streamlit. The app sends the question to GraphRAG, which tries Neo4j graph-guided retrieval, blends evidence with BM25 plus dense retrieval, and asks Gemini to write an evidence-grounded answer with citations. If the graph is not useful, hybrid retrieval becomes fallback.",
        "Reem": "My part is the evidence base. The system starts with PDFs and metadata. We split documents into chunks, and every chunk keeps document ID and page number. This is why final answers can cite real source pages.",
        "Salma": "My part is retrieval. BM25 finds exact words, dense retrieval finds similar meaning, and hybrid retrieval combines both. I evaluated retrieval and latency so the fallback path is measured, not guessed.",
        "Rana": "My part is the graph. The graph is in Neo4j and the graph query language is Cypher. The Cypher templates are in src/graph/cypher_queries.py. Python selects a template and executes it through the Neo4j driver.",
        "Aaya adaptation and tuning": "My part is online adaptation and QLoRA. Online adaptation updates retrieval behavior from feedback. QLoRA trains a small adapter on Kaggle GPU, not the whole model, and we save comparison outputs.",
        "Alia": "My part is integration and safety. The FastAPI search endpoint is in src/api/main.py and calls retriever.search. I also check faithfulness, relevance, citation correctness, and fallback behavior.",
    }
    doc.add_heading("Exact script by member", level=1)
    for k, v in scripts.items():
        doc.add_heading(k, level=2)
        para(doc, v)
    doc.add_heading("Best live demo questions", level=1)
    bullets(doc, [
        "What makes the FeederBW low-voltage grid dataset useful for energy-transition planning research?",
        "How can sowing-date adjustment help cereal crops respond to warming-driven phenology changes?",
        "Why do climate projections create vulnerability concerns for electrical substations and transformer loading practices?",
        "How could carbon-free Haber-Bosch ammonia production align with intermittent renewable energy?",
        "How could post-COVID green stimulus choices affect future warming trajectories?",
        "What does the literature say about gradient boosting methods for wind power forecasting?",
    ])
    doc.add_heading("What to point at in the app", level=1)
    bullets(doc, [
        "Input box: where the user query is created as a Python string.",
        "Left panel: GraphRAG trace with retrieval type, fallback, latency, Cypher template, graph hits, answer, citations.",
        "Right panel: hybrid BM25 plus dense retrieval comparison.",
        "Expected output box: saved audited result used to compare live behavior.",
    ])
    doc.add_heading("Doctor question quick answers", level=1)
    qs = [
        ("Where is the query created?", "app/streamlit_app.py, by st.text_input and example buttons."),
        ("What language is graph query?", "Cypher for Neo4j."),
        ("Where is Cypher?", "src/graph/cypher_queries.py."),
        ("Where is search API?", "src/api/main.py, POST /search."),
        ("Where does API call search?", "retriever.search(req.question, k=req.top_k, filters=req.filters)."),
        ("What if Gemini/Neo4j fails?", "Use hybrid fallback and validated saved CSV outputs."),
    ]
    for q, a in qs:
        doc.add_heading("Q: " + q, level=3)
        para(doc, "A: " + a)
    out = OUT / "Team_Streamlit_Demo_Speaking_Script.docx"
    doc.save(out)
    return out


def main():
    outs = []
    for name, cfg in MEMBERS.items():
        print("building", name)
        outs.append(member_doc(name, cfg))
    print("building master")
    outs.append(master_doc())
    print("building demo script")
    outs.append(demo_doc())
    print("DONE")
    for out in outs:
        print(out)


if __name__ == "__main__":
    main()
